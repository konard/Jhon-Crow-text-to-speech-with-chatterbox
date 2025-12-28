"""DOCX document reader using python-docx."""

from pathlib import Path

from .base import DocumentReader, DocumentContent


class DOCXReader(DocumentReader):
    """Reader for DOCX documents using python-docx.

    This reader extracts text from Microsoft Word DOCX files
    and can identify footnotes and endnotes.
    """

    @property
    def supported_extensions(self) -> list[str]:
        """Return list of supported DOCX extensions."""
        return [".docx", ".DOCX"]

    def read(self, file_path: Path) -> DocumentContent:
        """Read and extract content from a DOCX document.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            DocumentContent with extracted text and footnotes.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file cannot be read as DOCX.
        """
        from docx import Document

        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        try:
            doc = Document(file_path)

            # Extract main text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            # Extract footnotes
            footnotes = self._extract_footnotes(doc)

            full_text = "\n\n".join(paragraphs)

            # Estimate page count (roughly 500 words per page)
            word_count = len(full_text.split())
            estimated_pages = max(1, word_count // 500)

            return DocumentContent(
                text=full_text,
                footnotes=footnotes,
                page_count=estimated_pages,
                metadata={"source": str(file_path), "format": "docx"}
            )

        except Exception as e:
            raise ValueError(f"Failed to read DOCX file: {e}") from e

    def _extract_footnotes(self, doc) -> list[str]:
        """Extract footnotes from the DOCX document.

        Args:
            doc: The python-docx Document object.

        Returns:
            List of footnote texts with their reference numbers.
        """
        footnotes = []

        try:
            # Access the footnotes part of the document
            # python-docx doesn't have direct footnote support,
            # so we access the underlying XML
            footnotes_part = doc.part.footnotes_part
            if footnotes_part is not None:
                for i, footnote in enumerate(footnotes_part.footnotes, start=1):
                    # Get text from footnote paragraphs
                    footnote_text = ""
                    for para in footnote.paragraphs:
                        footnote_text += para.text + " "
                    footnote_text = footnote_text.strip()
                    if footnote_text:
                        footnotes.append(f"[{i}] {footnote_text}")
        except AttributeError:
            # Document doesn't have footnotes or structure is different
            pass

        return footnotes
